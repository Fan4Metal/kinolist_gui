import glob
import io
import logging
import os
import re
import sys
import textwrap
import time
from copy import deepcopy
from pathlib import Path

from requests import get
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from kinopoisk_unofficial.kinopoisk_api_client import KinopoiskApiClient
from kinopoisk_unofficial.request.films.film_request import FilmRequest
from kinopoisk_unofficial.request.staff.staff_request import StaffRequest
from mutagen.mp4 import MP4, MP4Cover, MP4StreamInfoError
from PIL import Image
from tqdm import tqdm
import PTN

LIB_VER = "0.2.19"

if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO,
                        format='[%(asctime)s]%(levelname)s:%(name)s:%(message)s',
                        datefmt='%d.%m.%Y %H:%M:%S')
log = logging.getLogger("Lib")


def is_api_ok(api):
    '''Проверка авторизации.'''
    try:
        api_client = KinopoiskApiClient(api)
        request = FilmRequest(328)
        api_client.films.send_film_request(request)
    except Exception:
        return False
    else:
        return True


def image_to_file(image):
    """Return `image` as PNG file-like object."""
    image_file = io.BytesIO()
    image.save(image_file, format="PNG")
    return image_file


def get_resource_path(relative_path):
    '''
    Определение пути для запуска из автономного exe файла.

    Pyinstaller cоздает временную папку, путь в _MEIPASS.
    '''
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


def copy_table_after(table, paragraph):
    '''Копирование таблицы в указанный параграф.'''
    tbl, p = table._tbl, paragraph._p
    new_tbl = deepcopy(tbl)
    p.addnext(new_tbl)


def clone_first_table(document: Document, num):
    '''Клонирует первую таблицу в документе num раз.'''
    template = document.tables[0]
    paragraph = document.paragraphs[0]
    for _i in range(num):
        copy_table_after(template, paragraph)
        paragraph = document.add_paragraph()


def find_kp_id_in_title(title: str):
    """Находит тег KP~xxx в названии и возвращает xxx (kinopoisk id)

    Args:
        title (str): Название фильма

    Returns:
        _type_: Значение xxx (kinopoisk id)
    """
    id = re.search(r"KP~(\d+)", title)
    if id:
        return id.group(1)


def find_kp_id(film_list: list, api: str):
    """Gets list of kinopoisk ids for list of films

    Args:
        film_list (list): List of movie titles for search
        api (string): Kinopoisk API token

    Returns:
        list: List of two elements:
                 0. list of found kinopoisk ids
                 1. list of items that have not been found
    """
    film_codes = []
    film_not_found = []
    for film in film_list:

        code_in_name = find_kp_id_in_title(film)
        if code_in_name:
            try:
                film_info = get_film_info(code_in_name, api)
                log.info(f'Найден фильм: {film_info[0]} ({film_info[1]}), kinopoisk id: {code_in_name}')
                film_codes.append(code_in_name)
                continue
            except Exception:
                film_not_found.append(code_in_name)
                continue
        time.sleep(0.2)
        payload = {'keyword': film, 'page': 1}
        headers = {'X-API-KEY': api, 'Content-Type': 'application/json'}
        try:
            r = get('https://kinopoiskapiunofficial.tech/api/v2.1/films/search-by-keyword',
                             headers=headers,
                             params=payload)
            if r.status_code == 200:
                resp_json = r.json()
                if resp_json['searchFilmsCountResult'] == 0:
                    log.info(f'{film} не найден')
                    film_not_found.append(film)
                    continue
                else:
                    id = resp_json['films'][0]['filmId']
                    year = resp_json['films'][0]['year']
                    if 'nameRu' in resp_json['films'][0]:
                        found_film = resp_json['films'][0]['nameRu']
                    else:
                        found_film = resp_json['films'][0]['nameEn']
                    log.info(f'Найден фильм: {found_film} ({year}), kinopoisk id: {id}')
                    film_codes.append(id)
            else:
                log.warning('Ошибка доступа к https://kinopoiskapiunofficial.tech')
                return
        except Exception as e:
            log.warning("Exeption:", str(e))
            log.info(f'{film} не найден (exeption)')
            film_not_found.append(film)
            continue
    return [film_codes, film_not_found]


def find_kp_id2(film: str, api: str):
    result = []
    code_in_name = find_kp_id_in_title(film)
    if code_in_name:
        try:
            film_info = get_film_info(code_in_name, api)
            log.info(f'Найден фильм: {film_info[0]} ({film_info[1]}), kinopoisk id: {code_in_name}')
            result.append(code_in_name)
            result.append(film_info[0])
            result.append({film_info[1]})
            return result
        except Exception:
            return result
    payload = {'keyword': film, 'page': 1}
    headers = {'X-API-KEY': api, 'Content-Type': 'application/json'}
    try:
        r = get('https://kinopoiskapiunofficial.tech/api/v2.1/films/search-by-keyword',
                         headers=headers,
                         params=payload)
        if r.status_code == 200:
            resp_json = r.json()
            if resp_json['searchFilmsCountResult'] == 0:
                log.info(f'{film} не найден')
                return result
            else:
                id = resp_json['films'][0]['filmId']
                year = resp_json['films'][0]['year']
                if 'nameRu' in resp_json['films'][0]:
                    found_film = resp_json['films'][0]['nameRu']
                else:
                    found_film = resp_json['films'][0]['nameEn']
                log.info(f'Найден фильм: {found_film} ({year}), kinopoisk id: {id}')
                result.append(id)
                result.append(found_film)
                result.append(year)
                return result
        else:
            log.warning('Ошибка доступа к https://kinopoiskapiunofficial.tech')
            return result
    except Exception as e:
        log.warning("Exeption:", str(e))
        log.info(f'{film} не найден (exeption)')
        return result


def get_film_info(film_code, api, shorten=False):
    '''
    Получение информации о фильме с помощью kinopoisk_api_client.

            Элементы списка:
                0 - название фильма на русском языке
                1 - год
                2 - рейтинг Кинопоиска
                3 - список стран
                4 - описание
                5 - ссылка на постер
                6 - ссылка на превью постера
                7 - список режиссеров
                8 - список актеров
                9 - Постер размером 360x540 в формате PIL.Image.Image
    '''
    api_client = KinopoiskApiClient(api)
    request_staff = StaffRequest(film_code)
    response_staff = api_client.staff.send_staff_request(request_staff)

    directors_list = []
    for item in response_staff.items:
        if item.profession_text == 'Режиссеры':
            if item.name_ru == '':
                directors_list.append(item.name_en)
            else:
                directors_list.append(item.name_ru)

    staff_list = []
    for item in response_staff.items:
        if len(staff_list) == 10:
            break
        if item.profession_text == 'Актеры':
            if item.name_ru == '':
                staff_list.append(item.name_en)
            else:
                staff_list.append(item.name_ru)

    request_film = FilmRequest(film_code)
    response_film = api_client.films.send_film_request(request_film)
    # с помощью регулярного выражения находим значение стран в кавычках ''
    countries = re.findall("'([^']*)'", str(response_film.film.countries))
    # имя файла
    if response_film.film.name_ru:
        film_name = response_film.film.name_ru
    else:
        film_name = response_film.film.name_original

    # Сокращение описания фильма
    if shorten:
        description = response_film.film.description.replace("\n\n", " ")
        description = textwrap.shorten(description,
                                       665,
                                       fix_sentence_endings=True,
                                       break_long_words=False,
                                       placeholder='...')
    else:
        description = response_film.film.description

    film_list = [
        film_name, response_film.film.year, response_film.film.rating_kinopoisk, countries, description,
        response_film.film.poster_url, response_film.film.poster_url_preview
    ]
    result = film_list
    result.append(directors_list)
    result.append(staff_list)
    # загрузка постера
    cover_url = response_film.film.poster_url
    cover = get(cover_url, stream=True)
    if cover.status_code == 200:
        cover.raw.decode_content = True
        image = Image.open(cover.raw)
        width, height = image.size
        # обрезка до соотношения сторон 1x1.5
        if width > (height / 1.5):
            image = image.crop((((width - height / 1.5) / 2), 0, ((width - height / 1.5) / 2) + height / 1.5, height))
        elif height > (1.5 * width):
            image = image.crop((0, ((height - width * 1.5) / 2), width, ((height + width * 1.5) / 2)))
        image.thumbnail((360, 540))
        rgb_image = image.convert('RGB')  # Fix "OSError: cannot write mode RGBA as JPEG"
        result.append(rgb_image)
    else:
        result.append("")
    return result


def get_full_film_list(film_codes: list, api: str, shorten=False):
    """Загружает информацию о фильмах

    Args:
        film_codes (list): Список kinopoisk_id фильмов
        api (str): Kinopoisk API token
        shorten (boolean): Option to shorten movie descriptions
    Returns:
        list: Список с полной информацией о фильмах для записи в таблицу.
    """
    full_films_list = []
    for film_code in tqdm(film_codes, desc="Загрузка информации...   "):
        try:
            film_info = get_film_info(film_code, api, shorten)
            full_films_list.append(film_info)
        except Exception as e:
            log.warning("Exeption:", str(e))
        else:
            continue
    return full_films_list


def write_film_to_table(current_table, filminfo: list):
    """Заполнение таблицы в файле docx.

    Args:
        current_table (Document object loaded from *docx*): указатель на текущую таблицу
        filminfo (list): информация о фильме
    """
    paragraph = current_table.cell(0, 1).paragraphs[0]  # название фильма + рейтинг
    if filminfo[2] is None:
        run = paragraph.add_run(str(filminfo[0]) + ' - ' + 'нет рейтинга')
    else:
        run = paragraph.add_run(str(filminfo[0]) + ' - ' + 'Кинопоиск ' + str(filminfo[2]))
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True

    paragraph = current_table.cell(1, 1).add_paragraph()  # год
    run = paragraph.add_run(str(filminfo[1]))
    run.font.name = 'Arial'
    run.font.size = Pt(10)

    paragraph = current_table.cell(1, 1).add_paragraph()  # страна
    run = paragraph.add_run(', '.join(filminfo[3]))
    run.font.name = 'Arial'
    run.font.size = Pt(10)

    paragraph = current_table.cell(1, 1).add_paragraph()  # режиссер
    if len(filminfo[7]) > 1:
        run = paragraph.add_run('Режиссеры: ' + ', '.join(filminfo[7]))
    elif filminfo[7]:
        run = paragraph.add_run('Режиссер: ' + filminfo[7][0])
    run.font.name = 'Arial'
    run.font.size = Pt(10)

    paragraph = current_table.cell(1, 1).add_paragraph()

    paragraph = current_table.cell(1, 1).add_paragraph()  # в главных ролях
    run = paragraph.add_run('В главных ролях: ')
    run.font.color.rgb = RGBColor(255, 102, 0)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run = paragraph.add_run(', '.join(filminfo[8]))
    run.font.color.rgb = RGBColor(0, 0, 255)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.underline = True

    paragraph = current_table.cell(1, 1).add_paragraph()
    paragraph = current_table.cell(1, 1).add_paragraph()
    paragraph = current_table.cell(1, 1).add_paragraph()  # синопсис
    run = paragraph.add_run(filminfo[4])
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    paragraph = current_table.cell(1, 1).add_paragraph()

    # запись постера в таблицу
    paragraph = current_table.cell(0, 0).paragraphs[1]
    run = paragraph.add_run()
    run.add_picture(image_to_file(filminfo[9]), width=Cm(7))


def write_all_films_to_docx(document, films: list, path: str):
    """Записывает информацию о фильмах в таблицы файла docx

    Args:
        document (_type_): Объект файла docx
        films (list): Список с информацией о фильмах
        path (str): Путь и имя для сохранения нового файла docx

    """
    table_num = len(films)
    if table_num > 1:
        clone_first_table(document, table_num - 1)
    for i in tqdm(range(table_num), desc="Запись в таблицу...      "):
        current_table = document.tables[i]
        write_film_to_table(current_table, films[i])
    try:
        document.save(path)
        log.info(f'Файл "{path}" создан.')
    except PermissionError:
        log.warning(f"Ошибка! Нет доступа к файлу {path}. Список не сохранен.")
        raise Exception


def file_to_list(file: str):
    """Читает текстовый файл и возвращает список строк

    Args:
        file (str): Текстовый файл

    Raises:
        FileNotFoundError: файл не найден

    Returns:
        list: список строк из файла
    """
    if os.path.isfile(file):
        with open(file, 'r', encoding="utf-8") as f:
            list = [x.rstrip() for x in f]
        return list
    else:
        print(f'Файл {file} не найден.')
        raise FileNotFoundError


def write_tags_to_mp4(film: list, file_path: str):
    """Запись тегов в файл mp4.

    Args:
        film (list): Информация о фильме
        file_path (str): Путь к файлу mp4
    """
    try:
        video = MP4(file_path)
    except MP4StreamInfoError as error:
        log.error(f"Ошибка! Не удалось открыть файл ({error}): {os.path.basename(file_path)}")
        return False
    video.delete()  # удаление всех тегов
    video["\xa9nam"] = film[0]  # title
    if film[4]:
        video["desc"] = film[4]  # description
        video["ldes"] = film[4]  # long description
    if film[1]:
        video["\xa9day"] = str(film[1])  # year
    video["covr"] = [MP4Cover(image_to_file(film[9]).getvalue(), imageformat=MP4Cover.FORMAT_PNG)]
    video.save()
    return True


def clear_tags(file_path: str):
    """Удаление тегов в файле mp4.

    Args:
        file_path (str): Путь к файлу mp4
    """
    try:
        video = MP4(file_path)
    except MP4StreamInfoError as error:
        log.error(f"Ошибка! Не удалось открыть файл ({error}): {os.path.basename(file_path)}")
        return False
    video.delete()  # удаление всех тегов
    video.save()
    return True


def docx_to_pdf_libre(file_in):
    file_in_abs = os.path.abspath(file_in)
    dir_out_abs = os.path.dirname(file_in_abs)
    soffice_path = "C:\Program Files\LibreOffice\program\soffice.exe"
    if not os.path.isfile(soffice_path):
        log.warning("Не найден файл soffice.exe. Возможно Libre Office не установлен.")
        return 1
    command = f'"{soffice_path}" --headless --convert-to pdf --outdir {dir_out_abs} {file_in_abs}'
    code_exit = os.system(command)
    return code_exit


def make_docx(kp_id_list: list, output: str, template: str, api: str, shorten: bool = False):
    file_path = get_resource_path(template)
    doc = Document(file_path)
    full_list = get_full_film_list(kp_id_list, api, shorten)
    write_all_films_to_docx(doc, full_list, output)


def rename_torrents(api: str, path=""):
    """Парсит имя из торрент файла и переименовывает в формат: название.ext

    Args:
        api (str): токен kinopoisk api
        path (str, optional): путь до файла. По умолчанию "".
    """
    files_paths = glob.glob(path)
    if len(files_paths) == 0:
        log.warning("Файлы не найдены.")
        return

    titles_all = []  # список кортежей (полный путь, имя файла, расширение)
    titles_parsed = []  # список распознанных парсером имен
    for file in files_paths:
        base_name = os.path.basename(file)
        tuple = (file,
                 os.path.splitext(base_name)[0],
                 os.path.splitext(base_name)[1])
        titles_all.append(tuple)
        titles_parsed.append(PTN.parse(base_name)['title'])

    id_list, films_not_found = find_kp_id(titles_parsed, api)
    if len(id_list) == 0:
        log.warning("Фильмы не найдены.")
        return

    titles_all_valid = []  # список кортежей (полный путь, имя файла, расширение) с только найденными фильмами
    for i in range(len(titles_all)):
        if titles_parsed[i] not in films_not_found:
            titles_all_valid.append(titles_all[i])

    films_info = get_full_film_list(id_list, api)

    for i, film in enumerate(films_info):
        trtable = film[0].maketrans('', '', '\/:*?"<>')
        file_name = film[0].translate(trtable)  # отфильтровываем запрещенные символы в новом имени файла
        os.rename(titles_all_valid[i][0],  # путь до исходного файла
                  os.path.join(
                        os.path.dirname(titles_all_valid[i][0]),  # путь до каталога исходного файла
                        f"{file_name}{titles_all_valid[i][2]}"  # новое имя файла + исходное расширение
                        )
                  )
        log.info(f"{titles_all_valid[i][1]}{titles_all_valid[i][2]} --> {file_name}{titles_all_valid[i][2]}")

    log.info("Файлы переименованы.")


def text_to_markdown(text: str) -> str:
    """Экранирует символы для вывода в режиме markdown

    Args:
        text (str): текст

    Returns:
        str: текст с экранированием
    """
    text_markdown = text.replace(".", "\.")
    text_markdown = text_markdown.replace("-", "\-")
    text_markdown = text_markdown.replace("(", "\(")
    text_markdown = text_markdown.replace(")", "\)")
    text_markdown = text_markdown.replace("!", "\!")
    return text_markdown


def main():
    import argparse_ru
    import argparse
    from config import KINOPOISK_API_TOKEN as api
    parser = argparse.ArgumentParser(prog='Kinolist_Lib',
                                     description='Библиотека для создания списков фильмов в формате docx.',
                                     formatter_class=argparse.RawDescriptionHelpFormatter,
                                     epilog="""
Примеры:
kl -m \"Terminator\" \"Terminator 2\" KP~319  --создает список list.docx из 3 фильмов: Terminator,
                                              Terminator 2 и Terminator 3 (*)
kl -f movies.txt -o movies.docx           --создает список movies.docx из всех фильмов в файле movies.txt
kl -t ./Terminator.mp4                    --записывает теги в файл Terminator.mp4 в текущем каталоге
kl -t c:\movies\Terminator.mp4            --записывает теги в файл Terminator.mp4 в каталоге c:\movies
kl --tag                                  --записывает теги во все mp4 файлы в текущем каталоге
kl --cleartags                            --удаляет все теги во всех mp4 файлах в текущем каталоге
kl -r *.mp4                               --переименовывает mp4 файлы в текущем каталоге (торрент -> название.mp4)
kl -l                                     --создает список list.docx из всех mp4 файлов в текущем каталоге

* Можно указать Kinopoisk_id напрямую, используя тег KP~XXX в названии фильма (где XXX - Kinopoisk_id)
                                        """)
    parser.add_argument("-ver",
                        "--version",
                        action="version",
                        version=f"%(prog)s {LIB_VER}",
                        help="выводит версию программы и завершает работу")
    parser.add_argument("-f",
                        "--file",
                        nargs=1,
                        help="создает список фильмов в формате docx из текстового файла в формате txt")
    parser.add_argument("-m", "--movie", nargs="+", help="создает список фильмов в формате docx из указанных фильмов")
    parser.add_argument("-o", "--output", nargs=1, help="имя выходного файла (list.docx по умолчанию)")
    parser.add_argument("-s",
                        "--shorten",
                        action='store_true',
                        help="сокращает описания фильмов, чтобы поместились два фильма на странице")
    parser.add_argument("-t",
                        "--tag",
                        nargs="?",
                        const=os.getcwd(),
                        help="записывает теги в файл mp4 (или во все mp4 файлы в текущем каталоге)")
    parser.add_argument("--cleartags",
                        nargs="?",
                        const=os.getcwd(),
                        help="удаляет все теги в файле mp4 (или во всех mp4 файлах в текущем каталоге)")
    parser.add_argument("-r", "--rename",
                        nargs="?",
                        const=os.getcwd(),
                        help="переименовывает mp4 файлыв в текущем каталоге")
    parser.add_argument("-l",
                        "--list",
                        nargs="?",
                        const=os.getcwd(),
                        help="создает список фильмов в формате docx из mp4 файлов в текущем каталоге")
    args = parser.parse_args()

    # определяем выходной файл
    if args.output:
        output = args.output[0]
        output_dir, output_file_name = os.path.split(output)
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        ext = os.path.splitext(output_file_name)[1]
        if ext != ".docx":
            print("Выходной файл должен иметь расширение docx.")
            return
    else:
        output = "list.docx"

    # список docx по списку txt
    if args.file:
        list = file_to_list((args.file[0]))
        if len(list) == 0:
            log.warning("Фильмы не найдены.")
            return
        log.info(f"Запрос из {args.file[0]} ({len(list)}): " + ", ".join(list))
        kp_codes = find_kp_id(list, api)
        if len(kp_codes[1]) != 0:
            for code in kp_codes[1]:
                log.warning(f"Фильм не найден, kinopoisk id: {code}")
        template = "template.docx"
        make_docx(kp_codes[0], output, template, api, args.shorten)

    # список docx из параметров
    elif args.movie:
        film = args.movie
        kp_codes = find_kp_id(film, api)
        if len(kp_codes[1]) != 0:
            for code in kp_codes[1]:
                log.warning(f"Фильм не найден, kinopoisk id: {code}")
        if len(kp_codes[0]) == 0:
            log.warning("Фильмы не найдены.")
            return
        template = "template.docx"
        make_docx(kp_codes[0], output, template, api, args.shorten)

    # запись тегов в mp4
    elif args.tag:
        path = args.tag
        if os.path.isfile(path):
            mp4_file = os.path.basename(path)
            name, ext = os.path.splitext(mp4_file)
            if ext != ".mp4":
                log.error("Можно записывать теги только в файлы mp4.")
                return
            name_list = []
            name_list.append(name)
            kp_ids = find_kp_id(name_list, api)
            if len(kp_ids[0]) == 0:
                log.warning("Фильм не найден.")
                return
            kp_id = kp_ids[0][0]
            film_info = get_film_info(kp_id, api)
            if not write_tags_to_mp4(film_info, path):
                log.warning(f"Тег не записан в файл: {mp4_file}")
                return
            log.info(f"Записан тег в файл: {mp4_file}")

        elif os.path.isdir(path):
            log.info(f"Поиск файлов mp4 в каталоге: {os.path.abspath(path)}")
            mp4_files = glob.glob(os.path.join(path, '*.mp4'))
            if len(mp4_files) == 0:
                log.warning(f'В каталоге "{path}" файлы mp4 не найдены.')
                return
            mp4_files_names = []
            for name in mp4_files:
                mp4_files_names.append(os.path.split(name)[1])
            for file in mp4_files_names:
                log.info(f"Найден файл: {file}")
            log.info(f"Всего найдено: {len(mp4_files)}")

            film_list = []
            for file in mp4_files:
                film_list.append(os.path.splitext(os.path.basename(file))[0])
            kp_ids, films_not_found = find_kp_id(film_list, api)
            mp4_files_valid = []
            for i in range(len(mp4_files)):
                if film_list[i] not in films_not_found:
                    mp4_files_valid.append(mp4_files[i])
            full_films_list = get_full_film_list(kp_ids, api)
            for i, film in enumerate(full_films_list):
                if not write_tags_to_mp4(film, mp4_files_valid[i]):
                    log.warning(f"Тег не записан в файл: {os.path.basename(mp4_files_valid[i])}")
                    return
                log.info(f"Записан тег в файл: {os.path.basename(mp4_files_valid[i])}")
        else:
            log.error("Неверно указан путь.")

    # очистка тегов
    elif args.cleartags:
        path = args.cleartags
        if os.path.isfile(path):
            mp4_file = os.path.basename(path)
            name, ext = os.path.splitext(mp4_file)
            if ext != ".mp4":
                log.error("Можно удалять теги только в файлах mp4.")
                return
            if not clear_tags(path):
                log.warning(f"Теги не удалены в файле: {os.path.basename(path)}")
                return
            log.info(f"Теги удалены в файле: {os.path.basename(path)}")

        elif os.path.isdir(path):
            log.info(f"Поиск файлов mp4 в каталоге: {os.path.abspath(path)}")
            mp4_files = glob.glob(os.path.join(path, '*.mp4'))
            if len(mp4_files) == 0:
                log.warning(f'В каталоге "{path}" файлы mp4 не найдены.')
                return
            for file in mp4_files:
                if not clear_tags(file):
                    log.warning(f"Ошибка! Теги не удалены в файле: {os.path.basename(file)}")
                    return
                log.info(f"Теги удалены в файле: {os.path.basename(file)}")
        else:
            log.error("Неверно указан путь.")

    # список по mp4 файлам
    elif args.list:
        path = args.list
        log.info(f"Поиск файлов mp4 в каталоге: {os.path.abspath(path)}")
        mp4_files = glob.glob(os.path.join(path, '*.mp4'))
        if len(mp4_files) == 0:
            log.warning(f'В каталоге "{path}" файлы mp4 не найдены.')
            return
        mp4_files_names = []
        for name in mp4_files:
            mp4_files_names.append(os.path.split(name)[1])
        for file in mp4_files_names:
            log.info(f"Найден файл: {file}")
        log.info(f"Всего файлов: {len(mp4_files)}")

        film_list = []
        for file in mp4_files:
            film_list.append(os.path.splitext(os.path.basename(file))[0])
        kp_id, films_not_found = find_kp_id(film_list, api)
        if len(films_not_found) > 0:
            log.warning("Следующие фильмы не найдены: " + ", ".join(films_not_found))
        template = "template.docx"
        make_docx(kp_id, output, template, api, args.shorten)

    # переимонование torrent файлов
    elif args.rename:
        rename_torrents(api, args.rename)


if __name__ == "__main__":
    main()
